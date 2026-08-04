"""Nạp văn bản luật từ nhiều định dạng file khác nhau về dạng text thuần."""

from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".html", ".htm"}


@dataclass
class LoadedDoc:
    file_name: str
    text: str
    warnings: list[str] = field(default_factory=list)
    n_pages: int = 0


# --------------------------------------------------------------- chuẩn hoá


def normalize_text(text: str) -> str:
    """Chuẩn hoá Unicode + dọn rác thường gặp khi bóc tách PDF văn bản luật."""
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace(" ", " ").replace("﻿", "")
    # Bỏ ký tự điều khiển lạ nhưng giữ xuống dòng và tab
    text = "".join(ch for ch in text if ch == "\n" or ch == "\t" or unicodedata.category(ch)[0] != "C")

    lines: list[str] = []
    for raw in text.split("\n"):
        line = re.sub(r"[ \t]+", " ", raw).strip()
        # Bỏ dòng chỉ có số trang, hoặc "Trang 3/45", hoặc chuỗi dấu chấm dẫn mục lục
        if re.fullmatch(r"(trang\s*)?\d{1,4}\s*(/\s*\d{1,4})?", line, flags=re.IGNORECASE):
            continue
        if re.fullmatch(r"[.\-–—_·•\s]{4,}", line):
            continue
        lines.append(line)

    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# --------------------------------------------------------------- các loader


def _load_pdf(path: Path, doc: LoadedDoc) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("Thiếu thư viện pypdf. Chạy: pip install pypdf") from exc

    reader = PdfReader(str(path))
    doc.n_pages = len(reader.pages)
    parts: list[str] = []
    empty_pages = 0
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:  # trang lỗi không nên làm hỏng cả file
            doc.warnings.append(f"Không đọc được 1 trang: {exc}")
            page_text = ""
        if len(page_text.strip()) < 20:
            empty_pages += 1
        parts.append(page_text)

    if doc.n_pages and empty_pages / doc.n_pages > 0.6:
        doc.warnings.append(
            "Phần lớn trang không bóc tách được chữ — file nhiều khả năng là PDF scan. "
            "Cần OCR trước khi nạp (ví dụ dùng ocrmypdf) thì RAG mới tra cứu được."
        )
    return "\n".join(parts)


def _load_docx(path: Path, doc: LoadedDoc) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("Thiếu thư viện python-docx. Chạy: pip install python-docx") from exc

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _load_html(path: Path, doc: LoadedDoc) -> str:
    raw = _read_text_file(path)
    raw = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", raw)
    raw = re.sub(r"(?i)<br\s*/?>", "\n", raw)
    raw = re.sub(r"(?i)</(p|div|tr|h[1-6]|li)>", "\n", raw)
    raw = re.sub(r"<[^>]+>", " ", raw)
    for entity, char in (("&nbsp;", " "), ("&amp;", "&"), ("&lt;", "<"), ("&gt;", ">"), ("&quot;", '"')):
        raw = raw.replace(entity, char)
    return raw


def _read_text_file(path: Path) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1258", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_bytes().decode("utf-8", errors="replace")


def load_document(path: Path) -> LoadedDoc:
    """Đọc 1 file văn bản luật, trả về text đã chuẩn hoá."""
    path = Path(path)
    doc = LoadedDoc(file_name=path.name, text="")
    ext = path.suffix.lower()

    if ext == ".pdf":
        raw = _load_pdf(path, doc)
    elif ext == ".docx":
        raw = _load_docx(path, doc)
    elif ext in {".html", ".htm"}:
        raw = _load_html(path, doc)
    elif ext in {".txt", ".md"}:
        raw = _read_text_file(path)
    elif ext == ".doc":
        raise RuntimeError(
            "Định dạng .doc (Word 97-2003) không đọc trực tiếp được. "
            "Hãy mở bằng Word và lưu lại thành .docx hoặc .pdf."
        )
    else:
        raise RuntimeError(f"Định dạng không hỗ trợ: {ext}. Hỗ trợ: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")

    doc.text = normalize_text(raw)
    if len(doc.text) < 100:
        doc.warnings.append("File có rất ít nội dung văn bản sau khi bóc tách.")
    return doc
